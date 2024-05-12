import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Mm
from docx.shared import Pt
import json
from openpyxl import load_workbook



# Card dimensions in mm
CARD_HEIGHT     = 36
CARD_WIDTH      = 80
CARDS_PER_PAGE  = 12 # NUmber of cards on a single page

COLOR_SUMMARY_TABLE_HEADER = "#f3f3f3"  # Light grey
TABLE_STYLE = "Table Grid"

# File names
INPUT_FILE_NAME     = 'Input_Data.xlsx'                 # Input file which shall contain 2 columns. The column "A" shall contain each text which goes on the front of the cards. Column "B" shall contain the items which goes on the back of the cards.
TEMPLATE_FILE_NAME  = 'Template.docx'                   # Template docx with pre-setted modifications. Use the original template to achieve proper results.
OUTPUT_FILE_NAME    = 'Printable_Flash_Cards.docx'      # The generated docs file will be named like this.

# Creating file pathes
WORKING_DIRECTORY   = os.getcwd()
INPUT_FILE_PATH     = os.path.join(WORKING_DIRECTORY,  INPUT_FILE_NAME)
TEMPLATE_FILE_PATH  = os.path.join(WORKING_DIRECTORY,  TEMPLATE_FILE_NAME)
OUTPUT_FILE_PATH    = os.path.join(WORKING_DIRECTORY,  OUTPUT_FILE_NAME)



#================================
# Python Docx related functions
#================================

def set_repeat_table_header(row):
    """Sets property 'repeat header row on every new page' of table.

    :param row: Header row
    :type row: [type]
    :return: Row
    :rtype: [type]
    """
    
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def shade_cells(cells, shade):
    """Gives background color to the inputted cells of a table in a .docx file.

    :param cells: The cells which you want to be colored. It could be a full row, such as the header row.
    :type cells: list
    :param shade: HEX color code like: "#f3f3f3" (grey)
    :type shade: str
    """
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)

def set_column_width(column, width):
    """Sets the width of the inputted column of a table in a .docx file.

    :param column: Column index
    :type column: [type]
    :param width: Desired width 
    :type width: [type]
    """
    for cell in column.cells:
        cell.width = width

def set_row_height(row, height):
    """Sets the width of the inputted column of a table in a .docx file.

    :param column: Column index
    :type column: [type]
    :param width: Desired width 
    :type width: [type]
    """
    for row in table.rows:
        row.height = height

def add_centered_row(row, content_of_cells):
    """Adds a new row to a table. Every cell will hold centered align text. 

    :param row: A list of cells of the row to be centered.
    :type row: list
    :param content_of_cells: The content of the centered row.
    :type content_of_cells: list
    """
    for cell in range(len(row)):
        paragraph = row[cell].paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        paragraph_format.space_before = Pt(3)
        paragraph_format.space_after = Pt(3)
        
        cell_run = paragraph.add_run(content_of_cells[cell])
        cell_run.bold = True

def add_lefty_row(row, content_of_cells):
    """Adds a new row to a table. Every cell will hold left aligned text. 

    :param row: A list of cells of the row to be left aligned.
    :type row: list
    :param content_of_cells: The content of the lefty row.
    :type content_of_cells: list
    """
    for cell in range(len(row)):
        paragraph = row[cell].paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        paragraph_format.space_before = Pt(3)
        paragraph_format.space_after = Pt(3)
        
        cell_run = paragraph.add_run(content_of_cells[cell])

def add_heading_row(row, content_of_cells):
    """Adds a heading row to a table. It basically sets a different style than the rest of the table has for the inputted row.

    :param row: A list of cells of the header row
    :type row: list
    :param content_of_cells: The content of the header row.
    :type content_of_cells: list
    """
    for cell in range(len(row)):
        paragraph = row[cell].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_run = paragraph.add_run(content_of_cells[cell])
        cell_run.bold = True

#================================
# Excel file processing
#================================

def process_excel_sheet(excel_file_path, sheet_index, heading=True):
    wb = load_workbook(filename=excel_file_path)
    sheet = wb.worksheets[sheet_index]

    rows_to_process = []
    heading_row = []
    return_dict = {}

    for row in sheet.iter_rows():
        row_values = [cell.value for cell in row]
        rows_to_process.append(row_values)

    if heading:
        heading_row = rows_to_process.pop(0)
    else:
        heading_row = ['Unnamed Property {}'.format(i) for i in range(len(rows_to_process[0]))]

    for idx, row in enumerate(rows_to_process, start=1):
        processed_row = {heading_row[i]: row[i] for i in range(len(row))}
        return_dict['row_{}'.format(idx)] = processed_row

    return return_dict


#================================
# Main functionality
#================================

def create_doc(input_dict):

    
    document = Document(TEMPLATE_FILE_PATH)
    
    input_dict_iterator = iter(input_dict)

    flash_card_front_content_list   = []
    flash_card_back_content_list    = []

    
    input_dict_last_key     = list(input_dict.keys())[-1]
    last_key_reached = False

    for item in range(0, len(input_dict)):

        flash_item = next(input_dict_iterator)
        flash_item = input_dict[flash_item]
        
        flash_item_values   = list(flash_item.values())

        flash_card_front_content_list.append(flash_item_values[0])
        flash_card_back_content_list.append(flash_item_values[1])

        if flash_item == input_dict_last_key:
            last_key_reached = True

        if last_key_reached:
            empty_cells_front   = CARDS_PER_PAGE - len(flash_card_front_content_list)
            empty_cells_back    = CARDS_PER_PAGE - len(flash_card_back_content_list)

            for i in range(0, empty_cells_front + 1):
                flash_card_front_content_list.append("-")
            
            for i in range(0, empty_cells_back + 1):
                flash_card_back_content_list.append("-")
            

        if (len(flash_card_front_content_list) == CARDS_PER_PAGE and len(flash_card_back_content_list) == CARDS_PER_PAGE):

            front_table = document.add_table(rows=0, cols=2, style=TABLE_STYLE)

            row_cntr = 0
            for front_row_item in range(0, len(flash_card_front_content_list)):         
                if (front_row_item % 2) == 0:
                    front_row = []
                    row = front_table.add_row()
                    front_row.append(flash_card_front_content_list[front_row_item])
                    front_row.append(flash_card_front_content_list[front_row_item + 1])
                    add_centered_row(front_table.rows[row_cntr].cells, front_row)
                    row_cntr += 1

            for column in front_table.columns:
                column.width = Mm(CARD_WIDTH)
            for row in front_table.rows:
                row.height = Mm(CARD_HEIGHT)
            document.add_page_break()
            document.add_paragraph("", style='Normal')

            back_table = document.add_table(rows=0, cols=2, style=TABLE_STYLE)

            row_cntr = 0
            for back_row_item in range(0, len(flash_card_back_content_list)):         
                if (back_row_item % 2) == 0:
                    back_row = []
                    row = back_table.add_row()
                    back_row.append(flash_card_back_content_list[back_row_item + 1])
                    back_row.append(flash_card_back_content_list[back_row_item])
                    add_centered_row(back_table.rows[row_cntr].cells, back_row)
                    row_cntr += 1

            for column in back_table.columns:
                column.width = Mm(CARD_WIDTH)
            for row in back_table.rows:
                row.height = Mm(CARD_HEIGHT)
            document.add_page_break()
            document.add_paragraph("", style='Normal')

            flash_card_front_content_list = []
            flash_card_back_content_list = []

        else:
            pass

    # ------
    # Add Table of Content like table to the end of the doc as a summary.
    # ------
    document.add_heading("Summary", level=1)
    document.add_paragraph("", style = 'Normal')

    summary_table = document.add_table(rows=1, cols=2, style=TABLE_STYLE)
    header_row = [
        'Front Item',
        'Back Item'
    ]
    add_heading_row(summary_table.rows[0].cells, header_row)

    # Giving background color to the first column cells.
    shade_cells([summary_table.cell(0, 0), summary_table.cell(0, 1)], COLOR_SUMMARY_TABLE_HEADER)

    row_cntr = 0
    for item in input_dict:
        new_row = []
        summary_table.add_row()
        single_item_dict = input_dict[item]
        keys    = list(single_item_dict.keys())
        values  = list(single_item_dict.values())

        new_row.append(values[0])
        new_row.append(values[1])

        row_cntr += 1

        add_lefty_row(summary_table.rows[row_cntr].cells, new_row)

    # Set column widths separately. Template page is portrait oriented. Sum width should be 16 cm
    set_column_width(summary_table.columns[0], Mm(40))
    set_column_width(summary_table.columns[1], Mm(120))

    # Repeat header row on pagebreak.
    set_repeat_table_header(summary_table.rows[0])

    document.save(OUTPUT_FILE_PATH)


def main():
    
    print("* Flash Card Generator *")
    print("Script started...")
    print("Processing input file: {}".format(INPUT_FILE_PATH))

    data = process_excel_sheet(INPUT_FILE_PATH, 0)

    print("Generating output document...")

    create_doc(data)

    print("Script finished.") 
    print("See the generated file here: {}".format(OUTPUT_FILE_PATH))

main()

